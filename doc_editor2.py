import streamlit as st
import io
import json
import re
import zipfile
from datetime import datetime
from typing import Dict, Any, List, Optional

# PDF
import pdfplumber
try:
    import fitz  # PyMuPDF
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

# DOCX
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# AI
import google.generativeai as genai
from openai import OpenAI

st.set_page_config(
    page_title="Professional CV Suite",
    page_icon="📄",
    layout="wide"
)


# ─────────────────────────────────────────────────────────────
#  HYPERLINK UTILITIES
# ─────────────────────────────────────────────────────────────

def add_hyperlink_to_paragraph(paragraph, display_text: str, url: str,
                                font_size_pt: float = None,
                                color: RGBColor = RGBColor(0x00, 0x56, 0xB3)):
    """
    Insert a genuine, clickable hyperlink run into *paragraph*.

    Builds the XML properly:
        <w:hyperlink r:id="rId5" w:history="1">
          <w:r>
            <w:rPr>
              <w:rStyle w:val="Hyperlink"/>
              <w:color w:val="0056B3"/>
              <w:u w:val="single"/>
              [<w:sz w:val="18"/>]
            </w:rPr>
            <w:t xml:space="preserve">display_text</w:t>
          </w:r>
        </w:hyperlink>

    This avoids the 'CT_Hyperlink has no attribute get_or_add_rPr' error
    caused by reparenting an already-bound CT_Run element.
    """
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")

    new_run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")

    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rPr.append(clr)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    if font_size_pt is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size_pt * 2)))
        rPr.append(sz)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = display_text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return new_run


# ─────────────────────────────────────────────────────────────
#  TEXT + HYPERLINK EXTRACTION
# ─────────────────────────────────────────────────────────────

def _nearest_word(cx: float, cy: float, words: list) -> str:
    """Return the text of the pdfplumber word whose centre is closest to (cx, cy)."""
    best, best_dist = "", float("inf")
    for w in words:
        wx = (w["x0"] + w["x1"]) / 2
        wy = (w["top"] + w["bottom"]) / 2
        d = (wx - cx) ** 2 + (wy - cy) ** 2
        if d < best_dist:
            best_dist = d
            best = w["text"]
    return best


def extract_text_from_pdf(raw_bytes: bytes) -> tuple[str, list]:
    """
    Extract text and hyperlinks from a PDF.

    Returns (full_text, links_found)
    links_found = [{"url": str, "text": str}, ...]

    Strategy:
      - pdfplumber  → text + word bounding boxes
      - PyMuPDF     → link rectangles + URIs
      - Match each link rect to nearest word, annotate text as [word](url)
    """
    links_found = []
    plumber_pages = []

    # ── pdfplumber pass ───────────────────────────────────────
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages:
            words = page.extract_words() or []
            text  = page.extract_text() or ""
            plumber_pages.append({
                "text": text,
                "words": words,
                "height": float(page.height),
            })

    # ── PyMuPDF pass ──────────────────────────────────────────
    if FITZ_AVAILABLE:
        doc = fitz.open(stream=raw_bytes, filetype="pdf")
        for page_num, fitz_page in enumerate(doc):
            if page_num >= len(plumber_pages):
                continue
            pd = plumber_pages[page_num]

            for link in fitz_page.get_links():
                if link.get("kind") != 2:   # 2 = URI
                    continue
                url = (link.get("uri") or "").strip()
                if not url or not url.startswith(("http", "mailto", "ftp")):
                    continue

                # fitz rect is top-left origin; pdfplumber is bottom-left origin
                rect = link["from"]
                cx = (rect.x0 + rect.x1) / 2
                cy_plumber = pd["height"] - (rect.y0 + rect.y1) / 2

                best_word = _nearest_word(cx, cy_plumber, pd["words"])
                link_text = best_word if best_word else url
                links_found.append({"url": url, "text": link_text})

                # Annotate the page text inline
                if best_word and best_word in pd["text"]:
                    pd["text"] = pd["text"].replace(
                        best_word, f"[{best_word}]({url})", 1
                    )
        doc.close()

    full_text = "\n\n".join(p["text"] for p in plumber_pages)
    return full_text, links_found


def extract_text_from_docx(cv_file) -> tuple[str, list]:
    """
    Extract text and hyperlinks from a DOCX file.

    Iterates each paragraph's XML children:
      - <w:hyperlink r:id="..."> → look up URL in part.rels, collect [text](url)
      - <w:r>                    → plain run text

    Returns (full_text, links_found)
    """
    doc = Document(cv_file)
    links_found = []
    text_parts  = []

    WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    REL_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    for para in doc.paragraphs:
        para_text = ""

        for child in para._p:
            tag = etree.QName(child).localname

            if tag == "hyperlink":
                r_id = (child.get(f"{{{REL_NS}}}id")
                        or child.get(qn("r:id")))
                url = ""
                if r_id:
                    rel = para.part.rels.get(r_id)
                    if rel:
                        url = rel.target_ref

                inner_text = "".join(
                    t.text or ""
                    for t in child.iter(f"{{{WORD_NS}}}t")
                )
                if url and inner_text:
                    links_found.append({"url": url, "text": inner_text})
                    para_text += f"[{inner_text}]({url})"
                else:
                    para_text += inner_text

            elif tag == "r":
                for t in child.iter(f"{{{WORD_NS}}}t"):
                    para_text += t.text or ""

        text_parts.append(para_text)

    return "\n".join(text_parts), links_found


# ─────────────────────────────────────────────────────────────
#  AI PROVIDER MANAGER
# ─────────────────────────────────────────────────────────────

class AIProvider:
    def __init__(self, provider: str, api_key: str, model: str = None):
        self.provider = provider
        self.api_key  = api_key
        self.models   = {
            "gemini":     "gemini-2.5-flash",
            "openrouter": "stepfun/step-3.5-flash:free",
        }
        self.model = model or self.models.get(provider, "gemini-2.5-flash")

        if provider == "gemini":
            genai.configure(api_key=api_key)
            self.gemini_model = genai.GenerativeModel(self.model)
        elif provider == "openrouter":
            self.client = OpenAI(
                base_url="https://openrouter.ai/api/v1",
                api_key=api_key,
                default_headers={
                    "HTTP-Referer": "https://cvsuite.app",
                    "X-Title": "CV Suite",
                },
            )

    def generate(self, prompt: str, temperature: float = 0.7) -> str:
        try:
            if self.provider == "gemini":
                response = self.gemini_model.generate_content(prompt)
                return response.text
            elif self.provider == "openrouter":
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=temperature,
                )
                return response.choices[0].message.content
        except Exception as e:
            st.error(f"AI Generation Error: {e}")
            return ""

    def get_available_models(self) -> List[str]:
        if self.provider == "openrouter":
            return [
                "stepfun/step-3.5-flash:free",
                "openai/gpt-4-turbo-preview",
                "openai/gpt-4",
                "anthropic/claude-3-opus",
                "anthropic/claude-3-sonnet",
                "google/gemini-pro",
                "mistralai/mistral-large",
                "meta-llama/llama-2-70b-chat",
            ]
        return ["gemini-2.5-flash", "gemini-1.5-pro"]


# ─────────────────────────────────────────────────────────────
#  CORE CV SUITE CLASS
# ─────────────────────────────────────────────────────────────

class CVSuite:
    def __init__(self, ai_provider: AIProvider):
        self.ai = ai_provider

    # ── TEXT EXTRACTION ───────────────────────────────────────

    def extract_text(self, cv_file) -> str:
        """Extract text and hyperlinks from PDF or DOCX, storing links in session state."""
        filename = cv_file.name.lower()

        if filename.endswith(".pdf"):
            if not FITZ_AVAILABLE:
                st.warning(
                    "📎 PyMuPDF not installed — PDF hyperlinks cannot be extracted. "
                    "Run: pip install PyMuPDF"
                )
            raw_bytes = cv_file.read()
            cv_file.seek(0)
            full_text, links_found = extract_text_from_pdf(raw_bytes)
            st.session_state.original_hyperlinks = links_found
            if links_found:
                st.success(f"🔗 Extracted {len(links_found)} hyperlink(s) from PDF.")
            else:
                st.info("📎 No hyperlinks detected in PDF. Use Word (.docx) for best link preservation.")
            return full_text

        else:  # docx
            full_text, links_found = extract_text_from_docx(cv_file)
            st.session_state.original_hyperlinks = links_found
            if links_found:
                st.success(f"🔗 Extracted {len(links_found)} hyperlink(s) from Word document.")
            return full_text

    # ── STEP 1: EXTRACT ALL CV DATA ───────────────────────────

    def extract_cv_data(self, cv_file) -> Dict:
        raw = self.extract_text(cv_file)
        prompt = f"""
You are a meticulous CV data extractor with expert-level attention to detail. Your job is to capture EVERY piece of information from this CV with 100% fidelity — do not skip, summarise, or fabricate anything.

IMPORTANT: Pay special attention to hyperlinks in the text. They appear in the format [text](url). Extract these links and preserve them in the appropriate fields (especially for projects, portfolio, GitHub, LinkedIn, etc.).

CV TEXT:
\"\"\"
{raw}
\"\"\"

Return ONLY a valid JSON object — no markdown fences, no explanation. Use this exact schema:

{{
  "personal_info": {{
    "name": "",
    "title": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin": "",
    "github": "",
    "portfolio": "",
    "website": "",
    "nationality": "",
    "date_of_birth": "",
    "id_number": "",
    "drivers_license": "",
    "gender": ""
  }},
  "professional_summary": "",
  "core_competencies": [],
  "skills": {{
    "technical": [],
    "soft": [],
    "tools": [],
    "languages": []
  }},
  "experience": [
    {{
      "title": "",
      "company": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "is_current": false,
      "summary": "",
      "responsibilities": [],
      "achievements": [],
      "technologies": [],
      "key_projects": []
    }}
  ],
  "education": [
    {{
      "degree": "",
      "field_of_study": "",
      "institution": "",
      "location": "",
      "start_date": "",
      "end_date": "",
      "grade": "",
      "achievements": [],
      "thesis": "",
      "relevant_modules": []
    }}
  ],
  "certifications": [
    {{
      "name": "",
      "issuer": "",
      "date": "",
      "expiry": "",
      "credential_id": ""
    }}
  ],
  "awards": [
    {{
      "name": "",
      "issuer": "",
      "date": "",
      "description": ""
    }}
  ],
  "volunteer_experience": [
    {{
      "role": "",
      "organization": "",
      "start_date": "",
      "end_date": "",
      "description": "",
      "achievements": []
    }}
  ],
  "publications": [
    {{
      "title": "",
      "publisher": "",
      "authors": "",
      "date": "",
      "link": "",
      "doi": ""
    }}
  ],
  "projects": [
    {{
      "name": "",
      "description": "",
      "technologies": [],
      "link": "",
      "dates": "",
      "outcomes": []
    }}
  ],
  "professional_memberships": [
    {{
      "organization": "",
      "role": "",
      "start_date": "",
      "end_date": ""
    }}
  ],
  "references": [
    {{
      "name": "",
      "title": "",
      "company": "",
      "email": "",
      "phone": "",
      "relationship": ""
    }}
  ],
  "interests": [],
  "additional_info": ""
}}

Critical rules for hyperlinks:
- Extract any hyperlinks found in the text (format: [text](url))
- For project links, store the URL in the "link" field
- For personal info (GitHub, LinkedIn, portfolio, website), extract the URLs
- If a project has a hyperlink, preserve it exactly
- For publications, extract DOI and links
- Do NOT lose any hyperlink information

Other rules:
- Extract dates EXACTLY as written in the CV
- Separate responsibilities (day-to-day duties) from achievements (results with impact)
- If the CV has a summary/objective, extract it verbatim into professional_summary
- Capture ALL skills mentioned anywhere in the document
- Do NOT add information that is not in the CV
- Use empty string "" for missing text, [] for missing arrays, false for booleans
"""
        return self._parse_json(self.ai.generate(prompt, temperature=0.1), self._empty_cv())

    # ── STEP 2: TAILOR CV TO JOB ──────────────────────────────

    def tailor_cv(self, cv_data: Dict, job_description: str, company_name: str) -> Dict:
        prompt = f"""
You are a senior executive CV writer and career coach with 20 years of experience placing candidates at top companies (FAANG, Fortune 500, high-growth startups). Your expertise is in crafting compelling, results-driven CVs that get interviews.

CRITICAL: PRESERVE ALL HYPERLINKS from the original CV data. Do not remove or modify project links, portfolio URLs, GitHub links, or any other URLs. Keep them exactly as they are in the data structure.

ORIGINAL CV DATA (JSON):
{json.dumps(cv_data, indent=2)}

TARGET JOB DESCRIPTION:
\"\"\"
{job_description}
\"\"\"

TARGET COMPANY: {company_name or "Not specified"}

Follow ALL of these advanced rewriting instructions carefully, while PRESERVING all hyperlinks:

1. PROFESSIONAL SUMMARY (most critical — hiring managers spend 6-8 seconds here):
   - Write 4-5 punchy, high-impact sentences specifically targeting this exact role
   - Open with: [X] years of experience in [domain] with expertise in [top 2 matching strengths]
   - Include 2-3 concrete achievements with specific metrics (numbers, percentages, scale)
   - Demonstrate understanding of the target company's needs and how you solve them
   - End with a compelling value proposition for THIS company
   - BANNED phrases: "results-driven", "dynamic", "passionate about", "team player", "hard worker", "go-getter", "detail-oriented", "self-starter", "proactive"
   - Use power words: "Spearheaded", "Architected", "Orchestrated", "Engineered", "Catalyzed", "Transformed"

2. EXPERIENCE (the most scrutinised section):
   - Rewrite EVERY achievement bullet using the CAR-L method: Challenge → Action → Result → Learning/Impact
   - Every bullet MUST start with a strong past-tense action verb
   - Add quantified metrics to EVERY bullet possible

3. PROJECTS (CRITICAL FOR HYPERLINKS):
   - KEEP ALL hyperlinks exactly as they appear in the original data
   - For each project, ensure the "link" field contains the original URL
   - Do not remove or modify project links

4. PERSONAL INFO: Keep ALL personal_info fields exactly as extracted — do not alter name, email, phone, location, or any contact details, INCLUDING hyperlinks (LinkedIn, GitHub, portfolio, website)

5. PUBLICATIONS: Preserve all DOIs, links, and publication URLs

Add these analysis fields to the returned JSON:
- "match_score": integer 1-10
- "key_changes": list of the 8-10 most important changes made
- "ats_keywords": list of 15-20 critical keywords from the JD now present in the CV
- "emphasized_skills": skills that were boosted/moved to prominence
- "deemphasized_skills": skills removed or minimised
- "missing_qualifications": things the JD requires that the candidate genuinely lacks
- "interview_talking_points": 5-6 specific, compelling things to highlight in the interview
- "strengths_amplified": top 3 strengths that were emphasized in the tailoring
- "improvement_areas": 2-3 areas where the CV could be further improved

Return ONLY valid JSON — same structure as input, plus the analysis fields.
No markdown fences, no explanation outside the JSON.
"""
        return self._parse_json(self.ai.generate(prompt, temperature=0.3), cv_data)

    # ── STEP 3: GENERATE COVER LETTER ─────────────────────────

    def generate_cover_letter(self, cv_data: Dict, job_description: str, company_name: str) -> str:
        pi       = cv_data.get("personal_info", {})
        exp      = cv_data.get("experience", [])
        skills   = cv_data.get("skills", {})
        projects = cv_data.get("projects", [])

        top_achievements = []
        for e in exp[:3]:
            top_achievements.extend(e.get("achievements", [])[:2])

        project_links = [p.get("link", "") for p in projects if p.get("link")]

        prompt = f"""
You are an expert cover letter writer who crafts compelling, highly personalised letters that get interviews at top companies. Write a cover letter that tells a compelling story and makes the hiring manager excited to interview this candidate.

CANDIDATE:
- Name: {pi.get('name', '')}
- Current Title: {pi.get('title', '')}
- Summary: {cv_data.get('professional_summary', '')[:200]}
- Top experiences: {json.dumps(exp[:3], indent=2)}
- Key skills: {json.dumps(skills, indent=2)}
- Top achievements: {json.dumps(top_achievements[:3], indent=2)}
- Portfolio/GitHub/Projects: {json.dumps(project_links[:3], indent=2)}
- ATS keywords to incorporate: {cv_data.get('ats_keywords', [])[:10]}

JOB DESCRIPTION:
\"\"\"
{job_description}
\"\"\"

COMPANY: {company_name or "the company"}

Writing requirements:
- Exact length: 380-430 words (strict)
- Start directly with "Dear Hiring Manager,"
- Use a professional, confident, warm tone

STRUCTURE:

Paragraph 1 (Opening Hook, 2-3 sentences):
- Do NOT start with "I am writing to apply for..." or "I am excited to apply..."
- Open with a bold, compelling statement about WHY this role or company specifically excites you
- Immediately establish your strongest qualification

Paragraph 2 (Best Achievement, 3-4 sentences):
- Tell your single most impressive and relevant achievement as a mini-story
- Use STAR method: Situation → Task → Action → Result
- Include measurable results

Paragraph 3 (Second Strength & Job Fit, 3-4 sentences):
- Address a specific, important requirement from the job description directly
- Provide a concrete example

Paragraph 4 (Cultural Fit & Company Research, 2-3 sentences):
- Show genuine, specific interest in THIS company
- Reference their product, mission, or values

Paragraph 5 (Close, 2 sentences):
- Confident, direct call to action

NEVER use these phrases:
- "I believe I would be a great fit"
- "I am a quick learner"
- "I am passionate about"
- "Please find my CV attached"
- "I look forward to hearing from you at your earliest convenience"
- "Thank you for your time and consideration"

Format: Write as plain paragraphs — no bullet points, no headers, no markdown
"""
        return self.ai.generate(prompt, temperature=0.7).strip()

    # ── STEP 4: BUILD CV WORD DOCUMENT ────────────────────────

    def build_cv_docx(self, cv_data: Dict) -> bytes:
        doc = Document()

        # Page setup
        for sec in doc.sections:
            sec.top_margin    = Cm(1.8)
            sec.bottom_margin = Cm(1.8)
            sec.left_margin   = Cm(2.0)
            sec.right_margin  = Cm(2.0)

        # Default style
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        pi = cv_data.get("personal_info", {})

        # ── NAME ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run((pi.get("name") or "YOUR NAME").upper())
        run.font.size  = Pt(28)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(31, 73, 125)

        # ── TITLE ──
        if pi.get("title"):
            tp = doc.add_paragraph()
            tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tp.paragraph_format.space_after = Pt(4)
            tr = tp.add_run(pi["title"])
            tr.font.size = Pt(14)
            tr.font.color.rgb = RGBColor(89, 89, 89)

        # ── CONTACT LINE WITH REAL HYPERLINKS ──
        self._add_contact_line(doc, pi)

        # ── RULE ──
        self._paragraph_border(doc, color="1F497D", size=18)

        # ── PROFESSIONAL SUMMARY ──
        if cv_data.get("professional_summary"):
            self._heading(doc, "PROFESSIONAL SUMMARY")
            sp = doc.add_paragraph(cv_data["professional_summary"])
            sp.paragraph_format.space_after   = Pt(8)
            sp.paragraph_format.line_spacing  = 1.15

        # ── CORE COMPETENCIES ──
        if cv_data.get("core_competencies"):
            self._heading(doc, "CORE COMPETENCIES")
            items = cv_data["core_competencies"]
            for i in range(0, len(items), 3):
                row_items = items[i:i + 3]
                rp = doc.add_paragraph()
                rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rp.paragraph_format.space_after = Pt(2)
                for j, item in enumerate(row_items):
                    if j > 0:
                        rp.add_run("  •  ")
                    rp.add_run(item)
                for run in rp.runs:
                    run.font.size = Pt(10)

        # ── SKILLS ──
        skills = cv_data.get("skills", {})
        if any(skills.get(k) for k in ["technical", "soft", "tools", "languages"]):
            self._heading(doc, "SKILLS")
            for label, key in [
                ("Technical Skills", "technical"),
                ("Tools & Software",  "tools"),
                ("Soft Skills",       "soft"),
                ("Languages",         "languages"),
            ]:
                items = skills.get(key, [])
                if items:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(3)
                    run = p.add_run(f"{label}: ")
                    run.bold = True
                    run.font.size = Pt(10)
                    p.add_run("  •  ".join(items)).font.size = Pt(10)

        # ── EXPERIENCE ──
        if cv_data.get("experience"):
            self._heading(doc, "PROFESSIONAL EXPERIENCE")
            for exp in cv_data["experience"]:
                self._add_experience(doc, exp)

        # ── PROJECTS WITH HYPERLINKS ──
        if cv_data.get("projects"):
            self._heading(doc, "KEY PROJECTS")
            for proj in cv_data["projects"]:
                self._add_project(doc, proj)

        # ── EDUCATION ──
        if cv_data.get("education"):
            self._heading(doc, "EDUCATION")
            for edu in cv_data["education"]:
                self._add_education(doc, edu)

        # ── CERTIFICATIONS ──
        if cv_data.get("certifications"):
            self._heading(doc, "CERTIFICATIONS & TRAINING")
            for cert in cv_data["certifications"]:
                cp = doc.add_paragraph()
                cp.paragraph_format.space_after = Pt(3)
                nr = cp.add_run(cert.get("name", ""))
                nr.bold = True
                details = []
                if cert.get("issuer"):
                    details.append(cert["issuer"])
                if cert.get("date"):
                    details.append(cert["date"])
                if cert.get("expiry"):
                    details.append(f"Expires: {cert['expiry']}")
                if details:
                    dr = cp.add_run(f"  |  {' | '.join(details)}")
                    dr.font.color.rgb = RGBColor(89, 89, 89)
                    dr.font.size = Pt(9)

        # ── PUBLICATIONS WITH LINKS ──
        if cv_data.get("publications"):
            self._heading(doc, "PUBLICATIONS")
            for pub in cv_data["publications"]:
                pp = doc.add_paragraph()
                pp.paragraph_format.space_after = Pt(3)
                nr = pp.add_run(pub.get("title", ""))
                nr.bold = True
                details = []
                if pub.get("authors"):
                    details.append(pub["authors"])
                if pub.get("publisher"):
                    details.append(pub["publisher"])
                if pub.get("date"):
                    details.append(pub["date"])
                if details:
                    pp.add_run(f"  |  {' | '.join(details)}").font.color.rgb = RGBColor(89, 89, 89)
                link = pub.get("link") or pub.get("doi")
                if link:
                    lp = doc.add_paragraph()
                    lp.paragraph_format.left_indent  = Pt(12)
                    lp.paragraph_format.space_after  = Pt(4)
                    add_hyperlink_to_paragraph(lp, link, link, font_size_pt=9)

        # ── AWARDS ──
        if cv_data.get("awards"):
            self._heading(doc, "AWARDS & RECOGNITION")
            for award in cv_data["awards"]:
                ap = doc.add_paragraph()
                ap.paragraph_format.space_after = Pt(3)
                nr = ap.add_run(award.get("name", ""))
                nr.bold = True
                details = []
                if award.get("issuer"):
                    details.append(award["issuer"])
                if award.get("date"):
                    details.append(award["date"])
                if details:
                    ap.add_run(f"  |  {' | '.join(details)}").font.color.rgb = RGBColor(89, 89, 89)
                if award.get("description"):
                    dp = doc.add_paragraph(award["description"])
                    dp.paragraph_format.left_indent = Pt(12)
                    dp.paragraph_format.space_after = Pt(4)

        # ── VOLUNTEER EXPERIENCE ──
        if cv_data.get("volunteer_experience"):
            self._heading(doc, "VOLUNTEER EXPERIENCE")
            for vol in cv_data["volunteer_experience"]:
                self._add_volunteer(doc, vol)

        # ── PROFESSIONAL MEMBERSHIPS ──
        if cv_data.get("professional_memberships"):
            self._heading(doc, "PROFESSIONAL MEMBERSHIPS")
            for mem in cv_data["professional_memberships"]:
                mp = doc.add_paragraph()
                mp.paragraph_format.space_after = Pt(3)
                nr = mp.add_run(mem.get("organization", ""))
                nr.bold = True
                if mem.get("role"):
                    mp.add_run(f"  |  {mem['role']}")
                dates = " – ".join(filter(None, [mem.get("start_date"), mem.get("end_date")]))
                if dates:
                    mp.add_run(f"  |  {dates}").font.color.rgb = RGBColor(120, 120, 120)

        # ── REFERENCES ──
        self._heading(doc, "REFERENCES")
        refs = cv_data.get("references", [])
        if refs:
            for ref in refs:
                rp = doc.add_paragraph()
                rp.paragraph_format.space_after = Pt(1)
                nr = rp.add_run(ref.get("name", ""))
                nr.bold = True
                if ref.get("title"):
                    rp.add_run(f",  {ref['title']}")
                if ref.get("company"):
                    rp.add_run(f"  |  {ref['company']}").font.color.rgb = RGBColor(89, 89, 89)
                contact = []
                if ref.get("email"):
                    contact.append(ref["email"])
                if ref.get("phone"):
                    contact.append(ref["phone"])
                if contact:
                    cp = doc.add_paragraph("  |  ".join(contact))
                    cp.paragraph_format.left_indent = Pt(12)
                    cp.paragraph_format.space_after = Pt(6)
        else:
            ip = doc.add_paragraph("Available upon request")
            ip.runs[0].italic = True
            ip.runs[0].font.color.rgb = RGBColor(120, 120, 120)

        # ── INTERESTS ──
        if cv_data.get("interests"):
            self._heading(doc, "INTERESTS")
            doc.add_paragraph("  •  ".join(cv_data["interests"]))

        # ── ADDITIONAL INFO ──
        if cv_data.get("additional_info"):
            self._heading(doc, "ADDITIONAL INFORMATION")
            doc.add_paragraph(cv_data["additional_info"])

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ── CONTACT LINE ──────────────────────────────────────────

    def _add_contact_line(self, doc, pi: dict):
        """Centred contact line — plain text for email/phone/location, real hyperlinks for URLs."""
        contact_fields = [
            ("email",     False),
            ("phone",     False),
            ("location",  False),
            ("linkedin",  True),
            ("github",    True),
            ("portfolio", True),
            ("website",   True),
        ]
        items = []
        for field, is_url in contact_fields:
            value = (pi.get(field) or "").strip()
            if not value:
                continue
            if is_url:
                display = (value
                           .replace("https://", "")
                           .replace("http://", "")
                           .replace("www.", ""))
                items.append((display, value))
            else:
                items.append((value, None))

        if not items:
            return

        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(2)

        for i, (display, url) in enumerate(items):
            if i > 0:
                sep = cp.add_run("  •  ")
                sep.font.size = Pt(9)
                sep.font.color.rgb = RGBColor(100, 100, 100)

            if url:
                add_hyperlink_to_paragraph(
                    cp, display, url,
                    font_size_pt=9,
                    color=RGBColor(100, 100, 100),
                )
            else:
                run = cp.add_run(display)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 100, 100)

    # ── PROJECT WITH HYPERLINK ────────────────────────────────

    def _add_project(self, doc, proj: dict):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(1)

        name_run = p.add_run(proj.get("name", ""))
        name_run.bold = True
        name_run.font.size = Pt(11)

        if proj.get("dates"):
            p.add_run(f"  |  {proj['dates']}").font.color.rgb = RGBColor(120, 120, 120)

        link = (proj.get("link") or "").strip()
        if link:
            p.add_run("  •  ")
            add_hyperlink_to_paragraph(p, "View Project", link, font_size_pt=9)

        if proj.get("description"):
            dp = doc.add_paragraph(proj["description"])
            dp.paragraph_format.left_indent = Pt(12)
            dp.paragraph_format.space_after = Pt(2)

        for outcome in proj.get("outcomes", []):
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent       = Pt(12)
            bp.paragraph_format.first_line_indent = Pt(-12)
            bp.paragraph_format.space_after        = Pt(2)
            bullet = bp.add_run("• ")
            bullet.font.color.rgb = RGBColor(31, 73, 125)
            bp.add_run(outcome).font.size = Pt(10)

        if proj.get("technologies"):
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent = Pt(12)
            tp.paragraph_format.space_after  = Pt(6)
            tr2 = tp.add_run("Technologies: ")
            tr2.bold = True
            tr2.font.size = Pt(9.5)
            tp.add_run(", ".join(proj["technologies"])).font.size = Pt(9.5)

    # ── EXPERIENCE BLOCK ──────────────────────────────────────

    def _add_experience(self, doc, exp: dict):
        start = exp.get("start_date", "")
        end   = "Present" if exp.get("is_current") else exp.get("end_date", "")
        dates = " – ".join(filter(None, [start, end]))

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(0)
        title_run = p.add_run(exp.get("title", ""))
        title_run.bold = True
        title_run.font.size = Pt(11)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        company_parts = []
        if exp.get("company"):
            company_parts.append(exp["company"])
        if exp.get("location"):
            company_parts.append(exp["location"])
        if company_parts:
            company_run = p2.add_run(" | ".join(company_parts))
            company_run.italic = True
            company_run.font.size = Pt(10)
            company_run.font.color.rgb = RGBColor(89, 89, 89)
        if dates:
            p2.add_run(f"  |  {dates}").font.color.rgb = RGBColor(120, 120, 120)

        if exp.get("summary"):
            sp = doc.add_paragraph(exp["summary"])
            sp.paragraph_format.left_indent = Pt(12)
            sp.paragraph_format.space_after = Pt(3)
            for r in sp.runs:
                r.italic = True
                r.font.size = Pt(9.5)

        for item in exp.get("responsibilities", []):
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent       = Pt(12)
            bp.paragraph_format.first_line_indent = Pt(-12)
            bp.paragraph_format.space_after        = Pt(2)
            bullet = bp.add_run("• ")
            bullet.font.color.rgb = RGBColor(31, 73, 125)
            bullet.font.size = Pt(10)
            bp.add_run(item).font.size = Pt(10)

        for item in exp.get("achievements", []):
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent       = Pt(12)
            bp.paragraph_format.first_line_indent = Pt(-12)
            bp.paragraph_format.space_after        = Pt(2)
            bullet = bp.add_run("• ")
            bullet.font.color.rgb = RGBColor(31, 73, 125)
            bullet.font.size = Pt(10)
            ar = bp.add_run(item)
            ar.font.size = Pt(10.5)  # achievements slightly larger

        if exp.get("technologies"):
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent = Pt(12)
            tp.paragraph_format.space_after  = Pt(6)
            tr2 = tp.add_run("Technologies: ")
            tr2.bold = True
            tr2.font.size = Pt(9.5)
            tp.add_run(", ".join(exp["technologies"])).font.size = Pt(9.5)

    # ── EDUCATION BLOCK ───────────────────────────────────────

    def _add_education(self, doc, edu: dict):
        degree = edu.get("degree", "")
        if edu.get("field_of_study"):
            degree = f"{degree} – {edu['field_of_study']}"
        dates = " – ".join(filter(None, [edu.get("start_date"), edu.get("end_date")]))

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(0)
        degree_run = p.add_run(degree)
        degree_run.bold = True
        degree_run.font.size = Pt(11)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        inst_parts = []
        if edu.get("institution"):
            inst_parts.append(edu["institution"])
        if edu.get("location"):
            inst_parts.append(edu["location"])
        if inst_parts:
            inst_run = p2.add_run(" | ".join(inst_parts))
            inst_run.italic = True
            inst_run.font.size = Pt(10)
            inst_run.font.color.rgb = RGBColor(89, 89, 89)
        if dates:
            p2.add_run(f"  |  {dates}").font.color.rgb = RGBColor(120, 120, 120)

        if edu.get("grade"):
            gp = doc.add_paragraph()
            gp.paragraph_format.left_indent = Pt(12)
            gp.paragraph_format.space_after = Pt(2)
            gp.add_run("Grade: ").bold = True
            gp.add_run(edu["grade"])

        if edu.get("thesis"):
            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent = Pt(12)
            tp.paragraph_format.space_after = Pt(2)
            tp.add_run("Thesis: ").bold = True
            tp.add_run(edu["thesis"])

        for item in edu.get("achievements", []):
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent       = Pt(12)
            bp.paragraph_format.first_line_indent = Pt(-12)
            bp.paragraph_format.space_after        = Pt(2)
            bullet = bp.add_run("• ")
            bullet.font.color.rgb = RGBColor(31, 73, 125)
            bp.add_run(item).font.size = Pt(10)

        if edu.get("relevant_modules"):
            mp = doc.add_paragraph()
            mp.paragraph_format.left_indent = Pt(12)
            mp.paragraph_format.space_after  = Pt(6)
            mp.add_run("Relevant modules: ").bold = True
            mp.add_run(", ".join(edu["relevant_modules"])).font.size = Pt(10)
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── VOLUNTEER BLOCK ───────────────────────────────────────

    def _add_volunteer(self, doc, vol: dict):
        dates = " – ".join(filter(None, [vol.get("start_date"), vol.get("end_date")]))

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(0)
        role_run = p.add_run(vol.get("role", ""))
        role_run.bold = True
        role_run.font.size = Pt(11)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        org_run = p2.add_run(vol.get("organization", ""))
        org_run.italic = True
        org_run.font.size = Pt(10)
        org_run.font.color.rgb = RGBColor(89, 89, 89)
        if dates:
            p2.add_run(f"  |  {dates}").font.color.rgb = RGBColor(120, 120, 120)

        if vol.get("description"):
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent       = Pt(12)
            bp.paragraph_format.first_line_indent = Pt(-12)
            bp.paragraph_format.space_after        = Pt(2)
            bullet = bp.add_run("• ")
            bullet.font.color.rgb = RGBColor(31, 73, 125)
            bp.add_run(vol["description"]).font.size = Pt(10)

        for item in vol.get("achievements", []):
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent       = Pt(12)
            bp.paragraph_format.first_line_indent = Pt(-12)
            bp.paragraph_format.space_after        = Pt(2)
            bullet = bp.add_run("• ")
            bullet.font.color.rgb = RGBColor(31, 73, 125)
            bp.add_run(item).font.size = Pt(10)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── DOCX STYLE HELPERS ────────────────────────────────────

    def _paragraph_border(self, doc, color="1F497D", size=12):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(8)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    str(size))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _heading(self, doc, text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
        run = p.add_run(text)
        run.font.name      = "Calibri"
        run.font.size      = Pt(12)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(31, 73, 125)
        run.font.all_caps  = True
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── COVER LETTER DOCUMENT ─────────────────────────────────

    def build_cover_letter_docx(self, text: str, pi: Dict, company: str = "") -> bytes:
        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Cm(2.5)
            sec.bottom_margin = Cm(2.5)
            sec.left_margin   = Cm(2.8)
            sec.right_margin  = Cm(2.8)

        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        # Name header
        name_p = doc.add_paragraph()
        name_p.paragraph_format.space_after = Pt(2)
        name_run = name_p.add_run((pi.get("name") or "").upper())
        name_run.font.size      = Pt(20)
        name_run.font.bold      = True
        name_run.font.color.rgb = RGBColor(31, 73, 125)

        # Contact info line (plain — cover letter doesn't need clickable links)
        contact = []
        if pi.get("email"):
            contact.append(pi["email"])
        if pi.get("phone"):
            contact.append(pi["phone"])
        if pi.get("location"):
            contact.append(pi["location"])
        if contact:
            contact_p = doc.add_paragraph("  |  ".join(contact))
            contact_p.paragraph_format.space_after = Pt(12)
            for run in contact_p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(100, 100, 100)

        self._paragraph_border(doc, color="1F497D", size=12)

        doc.add_paragraph(datetime.now().strftime("%d %B %Y")).paragraph_format.space_after = Pt(8)
        doc.add_paragraph("Hiring Manager").paragraph_format.space_after = Pt(2)
        if company:
            doc.add_paragraph(company).paragraph_format.space_after = Pt(12)

        for para in text.split("\n\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.paragraph_format.space_after  = Pt(10)
                p.paragraph_format.line_spacing = 1.15

        doc.add_paragraph()
        doc.add_paragraph("Sincerely,").paragraph_format.space_after = Pt(20)
        doc.add_paragraph(pi.get("name", ""))
        for field in ["email", "phone"]:
            if pi.get(field):
                doc.add_paragraph(pi[field])

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    # ── JSON HELPERS ──────────────────────────────────────────

    def _parse_json(self, raw: str, fallback: Dict) -> Dict:
        text = raw.strip()
        text = re.sub(r"^```json\s*", "", text)
        text = re.sub(r"^```\s*",     "", text)
        text = re.sub(r"\s*```$",     "", text)
        try:
            return json.loads(text)
        except Exception:
            match = re.search(r"\{.*\}", text, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group())
                except Exception:
                    pass
        return fallback

    def _empty_cv(self) -> Dict:
        return {
            "personal_info": {k: "" for k in [
                "name", "title", "email", "phone", "location", "linkedin",
                "github", "portfolio", "website", "nationality",
                "date_of_birth", "id_number", "drivers_license", "gender",
            ]},
            "professional_summary": "",
            "core_competencies": [],
            "skills": {"technical": [], "soft": [], "tools": [], "languages": []},
            "experience": [],
            "education": [],
            "certifications": [],
            "awards": [],
            "volunteer_experience": [],
            "publications": [],
            "projects": [],
            "professional_memberships": [],
            "references": [],
            "interests": [],
            "additional_info": "",
        }


# ─────────────────────────────────────────────────────────────
#  STREAMLIT APP
# ─────────────────────────────────────────────────────────────

def main():
    st.title("📄 Professional CV Suite")
    st.markdown("### AI-Powered CV Tailoring & Cover Letter Generation")

    # Initialise session state
    for key in ["cv_data", "tailored", "cover_letter", "cv_bytes", "cl_bytes",
                "job_desc_cache", "company_cache", "ai_provider", "original_hyperlinks"]:
        if key not in st.session_state:
            st.session_state[key] = None

    # ── SIDEBAR ──────────────────────────────────────────────
    with st.sidebar:
        st.header("⚙️ Configuration")

        provider = st.selectbox(
            "AI Provider",
            ["Gemini", "OpenRouter"],
            help="Choose your AI provider.",
        )

        if provider == "Gemini":
            api_key = st.text_input("Google Gemini API Key", type="password",
                                     help="Get your key at aistudio.google.com")
            model = st.selectbox("Model", ["gemini-2.5-flash", "gemini-1.5-pro"])
        else:
            api_key = st.text_input("OpenRouter API Key", type="password",
                                     help="Get your key at openrouter.ai")
            model = st.selectbox("Model", [
                "stepfun/step-3.5-flash:free",
                "openai/gpt-4-turbo-preview",
                "openai/gpt-4",
                "anthropic/claude-3-opus",
                "anthropic/claude-3-sonnet",
                "google/gemini-pro",
                "mistralai/mistral-large",
            ])

        if api_key:
            provider_key = "gemini" if provider == "Gemini" else "openrouter"
            st.session_state.ai_provider = AIProvider(provider_key, api_key, model)

        st.markdown("---")
        st.markdown("### 🔗 Hyperlink Support")
        st.markdown("""
- ✅ **PDF links** extracted via PyMuPDF
- ✅ **Word links** extracted from .rels XML
- ✅ **Clickable links** in output .docx files
- ✅ **Project / portfolio URLs** preserved
- ✅ **LinkedIn / GitHub** links preserved
        """)

        st.markdown("---")
        st.markdown("### 🚀 Features")
        st.markdown("""
- ✅ Multi-Provider AI: Gemini or OpenRouter
- ✅ Complete CV Extraction: 15+ sections
- ✅ Smart Tailoring: CAR method + metrics
- ✅ ATS Optimisation: Keyword embedding
- ✅ Hyperlink Preservation: All links kept
- ✅ Cover Letter Generation
- ✅ Match Scoring & Gap Analysis
        """)

        st.markdown("---")
        st.markdown("### 💡 Pro Tips")
        st.markdown("""
1. **Use Word (.docx)** for most reliable link extraction
2. **PDF links** work if they are proper annotations (not just typed URLs)
3. **Full JD** = better tailoring
4. **Review the Analysis tab** for interview prep
        """)

    # ── MAIN CONTENT ──────────────────────────────────────────
    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("📎 Upload Your CV")
        cv_file = st.file_uploader(
            "Choose your CV (PDF or Word)",
            type=["pdf", "docx"],
            help="Word format gives best hyperlink results",
        )
        if cv_file:
            st.success(f"✅ Uploaded: {cv_file.name}")

    with col2:
        st.subheader("💼 Job Details")
        company_name = st.text_input("Company Name (optional)",
                                      placeholder="e.g., Google, Microsoft…")
        job_description = st.text_area(
            "Job Description *",
            height=280,
            placeholder="Paste the full job description here…",
        )

    st.markdown("---")

    # ── GENERATE BUTTON ───────────────────────────────────────
    if st.button("🚀 Generate Tailored CV & Cover Letter", type="primary", use_container_width=True):
        if not api_key:
            st.error("❌ Please enter your API key in the sidebar.")
        elif not cv_file:
            st.error("❌ Please upload your CV.")
        elif not job_description.strip():
            st.error("❌ Please paste the job description.")
        elif not st.session_state.ai_provider:
            st.error("❌ Please configure the AI provider with a valid API key.")
        else:
            try:
                suite    = CVSuite(st.session_state.ai_provider)
                progress = st.progress(0)
                status   = st.empty()

                status.text("📄 Step 1/4 — Extracting all data from your CV (preserving hyperlinks)…")
                cv_data = suite.extract_cv_data(cv_file)
                st.session_state.cv_data = cv_data
                progress.progress(20)

                status.text("🎯 Step 2/4 — Tailoring CV to job description (keeping all hyperlinks)…")
                tailored = suite.tailor_cv(cv_data, job_description, company_name)
                st.session_state.tailored      = tailored
                st.session_state.job_desc_cache = job_description
                st.session_state.company_cache  = company_name
                progress.progress(55)

                status.text("✍️ Step 3/4 — Writing compelling cover letter…")
                cover = suite.generate_cover_letter(tailored, job_description, company_name)
                st.session_state.cover_letter = cover
                progress.progress(75)

                status.text("📝 Step 4/4 — Building professional Word documents with hyperlinks…")
                cv_bytes = suite.build_cv_docx(tailored)
                st.session_state.cv_bytes = cv_bytes
                cl_bytes = suite.build_cover_letter_docx(
                    cover, tailored.get("personal_info", {}), company_name
                )
                st.session_state.cl_bytes = cl_bytes
                progress.progress(100)

                status.text("✅ Complete!")
                st.success("✨ Your tailored CV and cover letter are ready — download below!")
                st.rerun()

            except Exception as e:
                st.error(f"Error: {str(e)}")
                import traceback
                st.code(traceback.format_exc())

    # ── DISPLAY RESULTS ───────────────────────────────────────
    if st.session_state.tailored and st.session_state.cv_bytes and st.session_state.cl_bytes:
        tailored = st.session_state.tailored
        cover    = st.session_state.cover_letter or ""
        pi       = tailored.get("personal_info", {})
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")

        project_links = [p.get("link", "") for p in tailored.get("projects", []) if p.get("link")]
        has_hyperlinks = bool(project_links or pi.get("linkedin") or pi.get("github") or pi.get("portfolio"))

        st.markdown("---")

        # Metrics
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("🎯 Match Score",       f"{tailored.get('match_score', '–')}/10")
        m2.metric("📊 Skills Highlighted", len(tailored.get("emphasized_skills", [])))
        m3.metric("🔑 ATS Keywords",       len(tailored.get("ats_keywords", [])))
        m4.metric("🔗 Project Links",      len(project_links), delta="Preserved")

        st.markdown("---")

        # Downloads
        st.subheader("📥 Download Your Documents")
        dl1, dl2, dl3 = st.columns(3)

        with dl1:
            st.markdown("**📄 Tailored CV**")
            st.caption("Rewritten & optimised for this job")
            st.download_button(
                label="⬇️ Download CV (.docx)",
                data=st.session_state.cv_bytes,
                file_name=f"tailored_cv_{ts}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_cv",
            )

        with dl2:
            st.markdown("**✉️ Cover Letter**")
            st.caption("Personalised for this role")
            st.download_button(
                label="⬇️ Download Cover Letter (.docx)",
                data=st.session_state.cl_bytes,
                file_name=f"cover_letter_{ts}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="dl_cl",
            )

        with dl3:
            st.markdown("**📦 Full Bundle**")
            st.caption("Both documents in ZIP")
            zip_buf = io.BytesIO()
            with zipfile.ZipFile(zip_buf, "w") as zf:
                zf.writestr(f"tailored_cv_{ts}.docx",     st.session_state.cv_bytes)
                zf.writestr(f"cover_letter_{ts}.docx", st.session_state.cl_bytes)
            zip_buf.seek(0)
            st.download_button(
                label="⬇️ Download ZIP Bundle",
                data=zip_buf.read(),
                file_name=f"application_bundle_{ts}.zip",
                mime="application/zip",
                use_container_width=True,
                key="dl_zip",
            )

        if has_hyperlinks:
            st.info("🔗 Your tailored CV contains clickable links — they will be active in the downloaded Word document.")

        st.markdown("---")

        # Tabs
        tab1, tab2, tab3, tab4 = st.tabs([
            "📄 CV Preview",
            "✉️ Cover Letter",
            "🎯 Analysis & ATS",
            "🔍 Raw Extracted Data",
        ])

        # ── TAB 1: CV PREVIEW ────────────────────────────────
        with tab1:
            st.subheader("Tailored CV Preview")

            with st.expander("👤 Personal Information", expanded=True):
                c1, c2 = st.columns(2)
                fields = [
                    ("Name",             "name"),
                    ("Title",            "title"),
                    ("Email",            "email"),
                    ("Phone",            "phone"),
                    ("Location",         "location"),
                    ("LinkedIn",         "linkedin"),
                    ("GitHub",           "github"),
                    ("Portfolio",        "portfolio"),
                    ("Website",          "website"),
                    ("Nationality",      "nationality"),
                    ("Date of Birth",    "date_of_birth"),
                    ("Driver's Licence", "drivers_license"),
                    ("ID Number",        "id_number"),
                    ("Gender",           "gender"),
                ]
                url_fields = {"linkedin", "github", "portfolio", "website"}
                for i, (label, key) in enumerate(fields):
                    val = pi.get(key, "")
                    if val:
                        col = c1 if i % 2 == 0 else c2
                        if key in url_fields:
                            col.markdown(f"**{label}:** [{val}]({val})")
                        else:
                            col.markdown(f"**{label}:** {val}")

            with st.expander("📝 Professional Summary", expanded=True):
                st.info(tailored.get("professional_summary", "—"))

            with st.expander("⭐ Core Competencies"):
                comps = tailored.get("core_competencies", [])
                cols  = st.columns(3)
                for i, comp in enumerate(comps):
                    cols[i % 3].markdown(f"• {comp}")

            with st.expander("🛠 Skills"):
                skills = tailored.get("skills", {})
                for label, key in [
                    ("💻 Technical Skills", "technical"),
                    ("🔧 Tools & Software",  "tools"),
                    ("🤝 Soft Skills",       "soft"),
                    ("🌍 Languages",         "languages"),
                ]:
                    if skills.get(key):
                        st.markdown(f"**{label}:** {', '.join(skills[key])}")

            with st.expander("💼 Professional Experience"):
                for exp in tailored.get("experience", []):
                    start = exp.get("start_date", "")
                    end   = "Present" if exp.get("is_current") else exp.get("end_date", "")
                    dates = " – ".join(filter(None, [start, end]))
                    st.markdown(f"### **{exp.get('title', '')}**")
                    st.markdown(f"*{exp.get('company', '')}* | {exp.get('location', '')} | {dates}")
                    for item in exp.get("responsibilities", []):
                        st.markdown(f"📌 {item}")
                    for item in exp.get("achievements", []):
                        st.markdown(f"🏆 {item}")
                    if exp.get("technologies"):
                        st.markdown(f"**Tech Stack:** {', '.join(exp['technologies'])}")
                    st.markdown("")

            with st.expander("🚀 Projects (with Links)"):
                for proj in tailored.get("projects", []):
                    st.markdown(f"### **{proj.get('name', '')}**")
                    if proj.get("link"):
                        st.markdown(f"🔗 [View Project]({proj['link']})")
                    if proj.get("description"):
                        st.markdown(f"*{proj['description']}*")
                    for outcome in proj.get("outcomes", []):
                        st.markdown(f"• {outcome}")
                    if proj.get("technologies"):
                        st.markdown(f"**Technologies:** {', '.join(proj['technologies'])}")
                    st.markdown("")

            with st.expander("🎓 Education"):
                for edu in tailored.get("education", []):
                    deg = edu.get("degree", "")
                    if edu.get("field_of_study"):
                        deg += f" – {edu['field_of_study']}"
                    st.markdown(f"**{deg}**")
                    st.markdown(f"*{edu.get('institution', '')}* | {edu.get('location', '')}")
                    if edu.get("start_date") or edu.get("end_date"):
                        st.markdown(f"{edu.get('start_date', '')} – {edu.get('end_date', '')}")
                    if edu.get("grade"):
                        st.markdown(f"**Grade:** {edu['grade']}")
                    if edu.get("thesis"):
                        st.markdown(f"**Thesis:** {edu['thesis']}")
                    for a in edu.get("achievements", []):
                        st.markdown(f"• {a}")
                    st.markdown("")

        # ── TAB 2: COVER LETTER ───────────────────────────────
        with tab2:
            st.subheader("Cover Letter")
            edited_cl = st.text_area(
                "Edit your cover letter (changes reflected in download):",
                value=cover,
                height=520,
                key="cl_edit",
            )
            col_cl1, col_cl2 = st.columns(2)
            with col_cl1:
                if st.session_state.ai_provider and edited_cl:
                    suite_cl    = CVSuite(st.session_state.ai_provider)
                    edited_bytes = suite_cl.build_cover_letter_docx(edited_cl, pi, company_name)
                    st.download_button(
                        label="⬇️ Download Edited Cover Letter (.docx)",
                        data=edited_bytes,
                        file_name=f"cover_letter_{ts}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="dl_cl_edited",
                    )
            with col_cl2:
                if st.button("🔄 Regenerate Cover Letter", use_container_width=True):
                    if st.session_state.ai_provider:
                        with st.spinner("Rewriting cover letter…"):
                            suite_regen = CVSuite(st.session_state.ai_provider)
                            new_cl = suite_regen.generate_cover_letter(
                                tailored,
                                st.session_state.job_desc_cache or "",
                                st.session_state.company_cache  or "",
                            )
                            st.session_state.cover_letter = new_cl
                            st.session_state.cl_bytes = suite_regen.build_cover_letter_docx(
                                new_cl, pi, company_name
                            )
                        st.rerun()

        # ── TAB 3: ANALYSIS ───────────────────────────────────
        with tab3:
            st.subheader("Analysis & Recommendations")

            if tailored.get("key_changes"):
                st.markdown("### 📝 Key Changes Made")
                for ch in tailored["key_changes"]:
                    st.markdown(f"✓ {ch}")
                st.markdown("---")

            col_a, col_b = st.columns(2)

            with col_a:
                if tailored.get("strengths_amplified"):
                    st.markdown("### 💪 Strengths Amplified")
                    for s in tailored["strengths_amplified"]:
                        st.markdown(f"✨ {s}")
                    st.markdown("")

                if tailored.get("ats_keywords"):
                    st.markdown("### 🔑 ATS Keywords Added")
                    kw_html = " ".join(
                        f'<span style="background:#e8f0fe;color:#1a73e8;padding:4px 10px;'
                        f'border-radius:20px;margin:4px;display:inline-block;font-size:0.85rem">'
                        f'{k}</span>'
                        for k in tailored["ats_keywords"][:15]
                    )
                    st.markdown(kw_html, unsafe_allow_html=True)

                if tailored.get("emphasized_skills"):
                    st.markdown("### ✅ Skills Emphasised")
                    for s in tailored["emphasized_skills"]:
                        st.markdown(f"↑ {s}")

            with col_b:
                if tailored.get("missing_qualifications"):
                    st.markdown("### ⚠️ Skill Gaps")
                    for q in tailored["missing_qualifications"]:
                        st.markdown(f"• {q}")

                if tailored.get("deemphasized_skills"):
                    st.markdown("### ↓ Skills De-emphasised")
                    for s in tailored["deemphasized_skills"]:
                        st.markdown(f"↓ {s}")

                if tailored.get("improvement_areas"):
                    st.markdown("### 📈 Areas for Improvement")
                    for area in tailored["improvement_areas"]:
                        st.markdown(f"💡 {area}")

            if tailored.get("interview_talking_points"):
                st.markdown("---")
                st.markdown("### 🎤 Interview Talking Points")
                for i, pt in enumerate(tailored["interview_talking_points"], 1):
                    st.markdown(f"**{i}.** {pt}")

            st.markdown("---")
            st.markdown("### 📊 CV Completeness Check")
            checks = {
                "Personal Information":    bool(pi.get("email") and pi.get("name")),
                "Professional Summary":    bool(tailored.get("professional_summary")),
                "Experience Section":      bool(tailored.get("experience")),
                "Education":               bool(tailored.get("education")),
                "Skills Section":          bool(any(tailored.get("skills", {}).values())),
                "Achievements with Metrics": any(
                    e.get("achievements") for e in tailored.get("experience", [])
                ),
                "Project Links":           bool(project_links),
                "ATS Keywords":            len(tailored.get("ats_keywords", [])) >= 10,
            }
            cc1, cc2 = st.columns(2)
            for i, (section, ok) in enumerate(checks.items()):
                (cc1 if i % 2 == 0 else cc2).markdown(f"{'✅' if ok else '⚠️ '} {section}")

        # ── TAB 4: RAW DATA ───────────────────────────────────
        with tab4:
            st.subheader("Extracted CV Data")
            st.caption("Original data extracted from your CV (before tailoring)")
            hl = st.session_state.original_hyperlinks or []
            if hl:
                st.info(f"🔗 Extracted {len(hl)} hyperlink(s) from your original CV")
                with st.expander("View extracted hyperlinks"):
                    for link in hl:
                        st.markdown(f"- **{link.get('text', '')}** → [{link.get('url', '')}]({link.get('url', '')})")
            st.json(st.session_state.cv_data or {})

    # Footer
    st.markdown("---")
    st.markdown(
        "<div style='text-align:center;color:gray;font-size:small'>"
        "Professional CV Suite · AI-Powered CV Tailoring & Cover Letter Generation<br>"
        "Powered by Gemini & OpenRouter AI · 🔗 Hyperlink preservation enabled"
        "</div>",
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    main()